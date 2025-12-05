import os
import platform

from typing import Tuple
from pydub import AudioSegment
from fpdf import FPDF
from docx import Document


def setup_ffmpeg() -> Tuple[str, str]:
    """
    Configura i percorsi degli eseguibili FFmpeg e FFprobe necessari per Pydub.

    Tenta di localizzare i binari nella directory del progetto per garantire
    che l'applicazione funzioni anche se FFmpeg non è installato a livello di sistema.

    Returns:
        Tuple[str, str]: Una tupla contenente i percorsi assoluti di (ffmpeg_path, ffprobe_path).
    """
    # Ottiene la directory corrente del file utils.py
    current_file_dir = os.path.dirname(os.path.abspath(__file__))

    # Risaliamo di un livello (assumendo che utils sia in src/ o src/utils/ e i binari siano nella root o src/)
    # NOTA: Adatta questo path se i tuoi exe sono in una posizione diversa.
    base_dir = os.path.dirname(current_file_dir)

    # Rilevamento del sistema operativo per l'estensione del file
    system_platform = platform.system()
    binary_ext = ".exe" if system_platform == "Windows" else ""

    # Costruzione dei percorsi assoluti
    ffmpeg_path = os.path.join(base_dir, f"ffmpeg{binary_ext}")
    ffprobe_path = os.path.join(base_dir, f"ffprobe{binary_ext}")

    # Aggiunge la directory dei binari al PATH di sistema temporaneo
    # Questo aiuta subprocess a trovare gli eseguibili senza percorsi completi se necessario
    os.environ["PATH"] += os.pathsep + base_dir

    # Configurazione esplicita di Pydub
    AudioSegment.converter = ffmpeg_path
    AudioSegment.ffmpeg = ffmpeg_path
    AudioSegment.ffprobe = ffprobe_path

    return ffmpeg_path, ffprobe_path


def get_transcripts_folder() -> str:
    """
    Restituisce il percorso della cartella centralizzata per le trascrizioni ('Sbobinature').
    Se la cartella non esiste, viene creata automaticamente.

    Returns:
        str: Il percorso assoluto della cartella di output.
    """
    # Ottiene la directory dove risiede questo script
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # Costruisce il percorso per la cartella di output
    folder_path = os.path.join(base_dir, "Sbobinature")

    # Pattern 'Look before you leap': crea solo se non c'è
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    return folder_path


def seconds_to_hms(seconds: float) -> str:
    """
    Converte un valore in secondi (float) in una stringa formattata HH:MM:SS.ss.
    Utile per visualizzare i timestamp nei log e nell'UI.

    Args:
        seconds (float): Il tempo in secondi.

    Returns:
        str: Stringa formattata (es. "01:05:30.50").
    """
    m, s = divmod(seconds, 60)
    h, m = divmod(m, 60)

    if h > 0:
        return "{:d}:{:02d}:{:05.2f}".format(int(h), int(m), s)
    else:
        return "{:02d}:{:05.2f}".format(int(m), s)


def hms_to_seconds(hms_str: str) -> float:
    """
    Converte una stringa formattata (HH:MM:SS o MM:SS) in secondi float.
    Utilizzato per interpretare l'input manuale dell'utente nelle Entry box.

    Args:
        hms_str (str): La stringa del tempo (es. "10:00").

    Returns:
        float: Il corrispettivo in secondi. Ritorna 0.0 se il formato è errato.
    """
    try:
        parts = hms_str.split(':')
        seconds = 0.0
        for part in parts:
            seconds = seconds * 60 + float(part)
        return seconds
    except ValueError:
        return 0.0


# --- EXPORT FUNCTIONS ---

def save_as_pdf(text: str, original_txt_path: str) -> str:
    """
    Esporta il testo trascritto in un file PDF.

    Nota Tecnica:
        FPDF standard non supporta nativamente Unicode completo (es. Emoji o caratteri speciali).
        Viene applicata una codifica 'latin-1' con replace per evitare crash durante la generazione.

    Args:
        text (str): Il contenuto della trascrizione.
        original_txt_path (str): Il percorso del file .txt originale (usato per derivare il nome).

    Returns:
        str: Il percorso del file PDF generato.
    """
    pdf = FPDF()
    pdf.add_page()

    # Impostazione Font per il Titolo
    pdf.set_font("Helvetica", size=12)  # Font base
    pdf.set_font("Helvetica", style="B", size=16)  # Grassetto

    # Titolo centrato
    filename = os.path.basename(original_txt_path)
    pdf.cell(0, 10, f"Trascrizione: {filename}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(10)  # Line break

    # Impostazione Font per il corpo del testo
    pdf.set_font("Helvetica", size=11)

    # Sanitizzazione del testo per FPDF (Fix per caratteri non supportati)
    safe_text = text.encode('latin-1', 'replace').decode('latin-1')

    # Scrittura del blocco di testo (multi_cell gestisce il word wrap automatico)
    pdf.multi_cell(0, 10, safe_text)

    # Output nella stessa cartella
    output_path = original_txt_path.replace(".txt", ".pdf")
    pdf.output(output_path)

    return output_path


def save_as_docx(text: str, original_txt_path: str) -> str:
    """
    Esporta il testo trascritto in un documento Word (.docx).
    Mantiene i paragrafi separati in base alle nuove righe nel testo originale.

    Args:
        text (str): Il contenuto della trascrizione.
        original_txt_path (str): Il percorso del file .txt originale.

    Returns:
        str: Il percorso del file .docx generato.
    """
    doc = Document()
    doc.add_heading(f'Trascrizione: {os.path.basename(original_txt_path)}', 0)

    # Aggiunge il testo paragrafo per paragrafo per mantenere la formattazione base
    for paragraph in text.split('\n'):
        if paragraph.strip():  # Ignora righe vuote
            doc.add_paragraph(paragraph)

    output_path = original_txt_path.replace(".txt", ".docx")
    doc.save(output_path)

    return output_path
